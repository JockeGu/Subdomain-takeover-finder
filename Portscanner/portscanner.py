"""
This is a portscanner.
"""
import socket
import pyfiglet
from docx import Document

BANNER = pyfiglet.figlet_format("Portscanner")
print(BANNER)

START_PORT = 1
END_PORT = 1001

document = Document()
document.add_heading("Open ports", 0)

with open('port_list.txt', mode='r', encoding='utf-8') as file:
    port_list = file.read().split(',')


def auto_scan(target, ports):
    """
    This function scans and prints out all open ports in a given list.
    :param target: Target IP-address
    :param ports: list of ports to scan.
    :return: returns open ports.
    """

    for port in ports:
        port = int(port)
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(0.2)
        result = sock.connect_ex((target, port))
        if result == 0:
            print(f"\033[1;32mPort {port} is open\033[0m")
            document.add_paragraph(f"Port {port} is open")
            document.save("open_ports.docx")
        sock.close()

    print("Scan finished.")


def basic_scan(target):
    """
    Scans and prints out all open ports from 1 to 1000
    :param target: Target IP-address
    :return: All open ports between 1-1000
    """
    open_ports = []

    for port in range(START_PORT, END_PORT + 1):
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(0.2)
        result = sock.connect_ex((target, port))
        if result == 0:
            open_ports.append(port)
            print(f"\033[1;32mPort {port} is open\033[0m")

        sock.close()

    print("Scan finished.")
    for port in open_ports:
        document.add_paragraph(f"Port {port} is open")

    document.save("open_ports.docx")


def main():
    """
    Lets the user select which mode to use. Auto or 1-1000 mode
    """
    target = str(input("Enter your target IP: "))
    valid_mode = ('1', '2')
    mode = ''

    while mode not in valid_mode:
        mode = str(input("Enter the mode you want to use (1, 2):"
                         "\n1. Auto(Popular ports)\n2. Ports 1-1000\n"))
        if mode == '1':
            print("Starting scan..")
            auto_scan(target, port_list)

        elif mode == '2':
            print("Starting scan..")
            basic_scan(target)
        else:
            print("Enter a valid option.")

    print("Result saved to open_ports.docx")


main()
